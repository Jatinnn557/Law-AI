from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, JsonResponse, HttpResponseRedirect
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from django.utils.timezone import now
from django.views.decorators.http import require_POST

from .forms import PDFUploadForm, RentalAgreementForm, DivorceAgreementForm, LandAgreementForm
from .models import UploadedPDF, ChatHistory
from .utils import process_pdf_and_summarize, extract_text_from_pdf

from bardapi import Bard
from docx import Document
from docxtpl import DocxTemplate
from docx2pdf import convert

import os
import uuid
import time
import pythoncom

from django.http import HttpResponseServerError, FileResponse
import tempfile
from io import BytesIO

from .forms import DocumentUploadForm
from .models import LegalDocument, Section
from .utils import classify_document, extract_text

from .forms import SectionForm
from django.views.decorators.csrf import csrf_exempt
from django.template.loader import render_to_string
from django.views.decorators.http import require_http_methods

from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.conf import settings
from django.core.mail import EmailMessage
from django.utils import timezone
from django.urls import reverse
from .models import *

@login_required
def home(request):
    return render(request, 'core/home.html')

def RegisterView(request):

    if request.method=="POST":
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        username = request.POST.get('username')
        email = request.POST.get('email')
        password = request.POST.get('password')

        user_data_has_error = False  #flag

        if User.objects.filter(username=username).exists():
            user_data_has_error = True
            messages.error(request,'username already exists')

        if User.objects.filter(email = email).exists():
            user_data_has_error = True
            messages.error(request, 'email already exists')
        
        if len(password)<8:
            user_data_has_error = True
            messages.error(request, 'password cannot be less than 8 characters')

        if user_data_has_error:
            return redirect('register')
        else:
            new_user = User.objects.create_user(
                first_name = first_name,
                last_name = last_name,
                email = email,
                username = username,
                password = password
            )
            messages.success(request, 'Account created Successfully. Login again to get started!')
            return redirect('login')


    return render(request, 'register.html')

def LoginView(request):

    if request.method =="POST":
        username = request.POST.get('username')
        password = request.POST.get('password')

        user = authenticate(request, username=username, password=password )
        
        if user is not None:
            login(request, user)
            return redirect('home')
        
        else: 
            messages.error(request, "Invalid Login Credentials!")
            return redirect('login')
        
    return render(request, 'login.html')

def LogoutView(request):
    logout(request)
    return redirect('login')

def ForgotPassword(request):

    if request.method=="POST":
        email=request.POST.get('email')

        try:
            user=User.objects.get(email=email)
            new_password_reset = PasswordReset(user=user)
            new_password_reset.save()

            password_reset_url = reverse('reset-password', kwargs={'reset_id': new_password_reset.reset_id})
            
            full_password_reset_url = f'{request.scheme}://{request.get_host()}{password_reset_url}'

            email_body = f'Reset your password using the link given below and DONT FORGET IT!!! :\n\n\n{full_password_reset_url}'
            
            email_message = EmailMessage(
                'Reset your password',
                email_body,
                settings.EMAIL_HOST_USER,
                [email]
            )

            email_message.fail_silently = True
            email_message.send() 

            return redirect(reverse('password-reset-sent', kwargs={'reset_id': new_password_reset.reset_id}))

        except User.DoesNotExist:
            messages.error(request, f"No User with email {email} found ")
            return redirect('forgot-password')

    return render (request, 'forgot_password.html')

def PasswordResetSent(request, reset_id):

    if PasswordReset.objects.filter(reset_id=reset_id).exists():
        return render(request, 'password_reset_sent.html')
    
    else:
        messages.error(request, 'Invalid reset id')
        return redirect('forgot-password')

def ResetPassword(request, reset_id):

    try:
        password_reset_id = PasswordReset.objects.get(reset_id=reset_id)

        if request.method == "POST":
            password = request.POST.get('password')
            confirm_password = request.POST.get('confirm_password')

            password_has_error = False

            if password != confirm_password:
                password_has_error = True
                messages.error(request, 'Passwords do not match! Please enter the same password in both fields.')

            
            if len(password) < 8:
                password_has_error = True
                messages.error(request, 'Password must be greater than 8 characters!')

            expiration_time = password_reset_id.created_when + timezone.timedelta(minutes=9)

            if timezone.now() > expiration_time:
                password_has_error = True
                messages.error(request, 'Reset Link has Expired!')
                password_reset_id.delete()
                return redirect('forgot-password')

            if not password_has_error:
                user= password_reset_id.user
                user.set_password(password)
                user.save()
                password_reset_id.delete()
                messages.success(request, 'Password reset Successfully!')
                return redirect ('login')

    except PasswordReset.DoesNotExist:
        messages.error(request, 'Invalid reset id')
        return redirect('forgot-password')
    
    return render(request, 'reset_password.html')

def upload_pdf(request):
    summary = None
    if request.method == 'POST':
        form = PDFUploadForm(request.POST, request.FILES)
        if form.is_valid():
            pdf = form.save()
            summary = process_pdf_and_summarize(pdf.file.path)
            pdf.summary = summary
            pdf.save()
    else:
        form = PDFUploadForm()
    return render(request, 'core/upload.html', {'form': form, 'summary': summary})


def chat_about_pdf(request, pdf_id):
    pdf = get_object_or_404(UploadedPDF, id=pdf_id)
    question = request.GET.get("q")
    answer = ""
    if question:
        full_text = extract_text_from_pdf(pdf.file.path)
        context = " ".join(full_text[:3])
        prompt = f"Based on this document:\n'''{context}'''\nAnswer this: {question}"
        try:
            answer = Bard().get_answer(prompt)['content']
        except:
            answer = "Could not process the question."

        # Save to chat history
        ChatHistory.objects.create(pdf=pdf, question=question, answer=answer)

    chat_history = pdf.chats.order_by('-asked_at')
    return render(request, 'core/chat.html', {
        'pdf': pdf,
        'answer': answer,
        'chat_history': chat_history
    })


def download_summary(request, pdf_id):
    pdf = get_object_or_404(UploadedPDF, id=pdf_id)
    response = HttpResponse(pdf.summary, content_type='text/plain')
    response['Content-Disposition'] = 'attachment; filename=summary.txt'
    return response


def document_selector(request):
    return render(request, 'law_documents/document_selector.html')


def generate_doc(request, form_class, template_name, doc_name, html_template):
    if request.method == 'POST':
        form = form_class(request.POST)
        if form.is_valid():
            doc_template_path = os.path.join(settings.BASE_DIR, 'core', 'docx_templates', template_name)
            doc = DocxTemplate(doc_template_path)
            doc.render(form.cleaned_data)

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{doc_name}.docx"'
            doc.save(response)
            return response
    else:
        form = form_class()
    return render(request, f'law_documents/{html_template}', {'form': form})


def generate_rental_agreement(request):
    return generate_doc(request, RentalAgreementForm, 'rental_template.docx', 'Rental_Agreement', 'generate_rental.html')


def generate_divorce_agreement(request):
    return generate_doc(request, DivorceAgreementForm, 'divorce_template.docx', 'Divorce_Agreement', 'generate_divorce.html')


def generate_land_agreement(request):
    return generate_doc(request, LandAgreementForm, 'land_template.docx', 'Land_Agreement', 'generate_land.html')


@require_POST
def preview_rental_doc(request):
    form = RentalAgreementForm(request.POST)
    if not form.is_valid():
        # Return errors as JSON or simple message
        return JsonResponse({'error': 'Invalid data submitted.'}, status=400)

    doc_template_path = os.path.join(settings.BASE_DIR, 'core', 'docx_templates', 'rental_template.docx')
    if not os.path.exists(doc_template_path):
        return JsonResponse({'error': 'Template not found.'}, status=404)

    # Load template and render with form data
    doc = DocxTemplate(doc_template_path)
    doc.render(form.cleaned_data)

    # Save rendered doc to in-memory file
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    # Load rendered doc to extract text
    document = Document(doc_io)
    content = []

    # Extract paragraphs
    for para in document.paragraphs:
        if para.text.strip():
            content.append(para.text)

    # Extract text from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    content.append(cell_text)

    # Join paragraphs with <br> for HTML display
    rendered_html = "<br>".join(content)

    # Return JSON with HTML content
    return JsonResponse({'rendered_html': rendered_html})



@require_POST
def preview_divorce_doc(request):
    form = DivorceAgreementForm(request.POST)
    if not form.is_valid():
        print("Form errors:", form.errors)  # Add this for debugging
        return JsonResponse({'error': 'Invalid data submitted.'}, status=400)

    doc_template_path = os.path.join(settings.BASE_DIR, 'core', 'docx_templates', 'divorce_template.docx')
    if not os.path.exists(doc_template_path):
        return JsonResponse({'error': 'Template not found.'}, status=404)

    doc = DocxTemplate(doc_template_path)
    doc.render(form.cleaned_data)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    document = Document(doc_io)
    content = []

    for para in document.paragraphs:
        if para.text.strip():
            content.append(para.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    content.append(cell_text)

    rendered_html = "<br>".join(content)
    return JsonResponse({'rendered_html': rendered_html})

@require_POST
def preview_land_doc(request):
    form = LandAgreementForm(request.POST)
    if not form.is_valid():
        return JsonResponse({'error': 'Invalid data submitted.'}, status=400)

    template_path = os.path.join(settings.BASE_DIR, 'core', 'docx_templates', 'land_template.docx')
    if not os.path.exists(template_path):
        return JsonResponse({'error': 'Template not found.'}, status=404)

    doc = DocxTemplate(template_path)
    doc.render(form.cleaned_data)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    document = Document(buffer)
    content = []

    for para in document.paragraphs:
        if para.text.strip():
            content.append(para.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    content.append(cell_text)

    return JsonResponse({'rendered_html': "<br>".join(content)})

@require_http_methods(["GET", "POST"])
def document_list(request):
    if request.method == 'POST':
        new_section_name = request.POST.get('new_section_name')
        if new_section_name:
            Section.objects.get_or_create(name=new_section_name)
            return redirect('document_list')

    sections = Section.objects.prefetch_related('documents').all()
    return render(request, 'core/document_list.html', {'sections': sections})

def edit_section(request, pk):
    section = get_object_or_404(Section, pk=pk)
    if request.method == 'POST':
        form = SectionForm(request.POST, instance=section)
        if form.is_valid():
            form.save()
            return redirect('document_list')
    else:
        form = SectionForm(instance=section)
    return render(request, 'core/edit_section.html', {'form': form, 'section': section})

@require_POST
def upload_to_section(request, section_id):
    section = get_object_or_404(Section, id=section_id)
    form = DocumentUploadForm(request.POST, request.FILES)
    if form.is_valid():
        uploaded_file = request.FILES['file']
        document = LegalDocument.objects.create(
            title=uploaded_file.name,
            section=section,
            file=uploaded_file
        )
        html = render_to_string('partials/document_item.html', {'doc': document})
        return JsonResponse({'success': True, 'html': html})
    return JsonResponse({'success': False, 'error': 'Invalid form'}, status=400)

def delete_document(request, doc_id):
    doc = get_object_or_404(LegalDocument, id=doc_id)
    doc.delete()
    return redirect('document_list')

def create_section(request):
    if request.method == 'POST':
        form = SectionForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('document_list')
    else:
        form = SectionForm()
    return render(request, 'core/create_section.html', {'form': form})

@require_http_methods(["GET", "POST"])
def auto_classify_upload(request):
    if request.method == 'POST':
        form = DocumentUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['file']
            text = extract_text(uploaded_file)
            section_name = classify_document(text)
            section, _ = Section.objects.get_or_create(name=section_name)
            LegalDocument.objects.create(
                title=uploaded_file.name,
                section=section,
                file=uploaded_file
            )
            return redirect('document_list')
    else:
        form = DocumentUploadForm()

    return render(request, 'core/upload2.html', {'form': form})